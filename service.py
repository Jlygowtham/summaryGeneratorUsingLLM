from db_connection import postgresDb
import fitz
from io import BytesIO
from docx import Document
import pandas as pd
from langchain_core.prompts import ChatPromptTemplate
from langchain_ollama.chat_models import ChatOllama
from fpdf import FPDF
from docx2pdf import convert
import os

def newUserService(userData: dict):
    try:
        name=userData.name
        email=userData.email
        password= userData.password

        pg= postgresDb()
        
        query="select email from userbio where email=%s"
        existUser= pg.showData(query,(email,))

        print(f"exist user status: {existUser}")

        if existUser.get('data'):
            return "user already exist"

        query="""
                Insert into UserBio(name,email,password)
                values(%s,%s,%s)
            """
        values=(name,email,password)

        insertNewUser = pg.insertData(query,values)
        
        print("New user status",insertNewUser['data'])
        return insertNewUser['data']
    
    except Exception as e:
        print(f"Error in newUserService function: {e}")
        return {"error":e,"status code":400}
    
    finally:
        pg.disconnect()

def loginService(userData:dict):
    try:
        email= userData.email
        password = userData.password

        pg= postgresDb()
        
        query="""
               Select * from userBio
               where email=%s and password=%s
         """
        values=(email,password)

        loginUser=pg.showData(query,values)

        print("login user status: ",loginUser)

        if not loginUser.get('data'):
            return "Invalid credentials"

        return loginUser

    except Exception as e:
        print(f"Error in loginService function: {e}")
        return {"error":e,"status code":400}
    
    finally:
        pg.disconnect()


def summaryService(file,file_type,user_id):
    try:
        if file_type=='pdf':
            data = pdfService(user_id,file,file_type)
            print("pdf summary:",data)
            return data
        
        elif file_type=='docx':
            data = docService(user_id,file,file_type)
            print("Document summary: ",data)
            return data
        
        elif file_type=='csv':
            csv_read= pd.read_csv(BytesIO(file.read()))
            content = csv_read.to_string()
            response = llmService(user_id,content,'csv')
            print("csv summary",response)
            insertLLMSummary(user_id,content,response,file_type)
            return response
        
        elif file_type=='xlsx':
            data = pd.read_excel(BytesIO(file.read()))
            content = data.to_string()
            response = llmService(user_id,content,'csv')
            print("xlsx summary",response)
            insertLLMSummary(user_id,content,response,file_type)
            return response

        else:
            data= file.read().decode('utf-8')
            response = llmService(user_id,data,'txt')
            print("text summary",response)
            insertLLMSummary(user_id,data,response,file_type)
            return response
        
    except Exception as e:
        print("Exception:",e)
        return {"error",e}
    

def pdfService(user_id,file,file_type):
    try:
       print("Enter pdf service")
       content=file.read()
       pdf=fitz.open(stream=content,filetype='pdf')
       txt=''
       for page in pdf:
           print("number: ",page.number)
           txt+=page.get_text()
       print("pdf content: ",txt)

       response = llmService(user_id,txt,'pdf')
       insertLLMSummary(user_id,txt,response,file_type)
       return response

    except Exception as e:
        print("Exception",e)
        return {'error': e}
    

def docService(userid,file,file_type):
    try:
        print("Enter document service")
        docRead=Document(BytesIO(file.read()))
        content = ''.join([para.text for para in docRead.paragraphs])
        response = llmService(userid,content,'docx')
        insertLLMSummary(userid,content,response,file_type)
        return response
    except Exception as e:
        print("Exception: ",e)
        return {"error":e}


def llmService(userid,content,file_type):
    try:
        model = ChatOllama(model='llama3.2:latest',temperature=0.3)

        prompt = ChatPromptTemplate.from_messages([
            ("system", f"""
Role:             
-You are a professional English teacher with over 10 years of experience and strong summarization skills.

Objective:
-Your job is to clearly and concisely summarize content based on its file type.

File Type: 
-{file_type if file_type else 'pdf or docx or csv or xlsx or txt file'}

Instructions:
- You must always produce a clear and meaningful summary based on the input — even if the content is informal, unstructured, or minimal.
- Do not generate SQL queries, programming code, shell commands, or advice.
- Do not include anything that wasn’t directly found or implied in the original content.
- Avoid hallucinating facts or adding hypothetical content.

Formatting:
        - Format your output using basic Markdown syntax that will be converted to a Word document.
        - Use only the following Markdown elements:
        - #, ##, ### for headings and subheadings
        - - for bullet points
        - **bold text** to emphasize key terms or sections
        - Do not use any other Markdown (like links, tables, or images).
        - Keep the summary clear and concise, following this structure:
            - One main heading for the document title
            - Section headings (## or ###) to group topics
            - Bullet points under each section with optional bold keywords

"""),
            ("user", "{content}")
        ])

        chain = prompt | model
        response = chain.invoke({"content": content})
        
        docxFilePath=convertSummaryToDocx(str(userid),response.content)
        pdfFilePath=f"{userid}_summary_pdf.pdf"
        convertSummaryToPdf(docxFilePath,pdfFilePath)
        print("LLm response:",response)
        return response.content

    except Exception as e:
        print("Exception: ",e)
        return {"error":e}
    

def insertLLMSummary(userid,content,llmresponse,file_type):
    try:
        query="""Insert into summaryHistory(fileType,content,llmSummary,user_id)
                 values(%s,%s,%s,%s)"""
        values=(file_type,content,llmresponse,str(userid))
        pg=postgresDb()
        result=pg.insertData(query,values)
        print("Result: ",result)

    except Exception as e:
        print("Exception",e)

    finally:
        pg.disconnect()


def convertSummaryToPdf(input_path,output_dir):
    try:
        if not os.path.exists(input_path):
            print("There is not path for document. So I can't convert docx to pdf")
            return

        # if output_dir and not os.path.exists(output_dir):
        #     os.makedirs(output_dir)

        convert(input_path)
        
        print("✅ Summary converted document to PDF successfully.")
    except Exception as e:
        print("Exception:",e)

def convertSummaryToDocx(userId,response):
    try:
        doc = Document()

        doc.add_heading("Summary", level=1)

        for line in response.splitlines():
            line = line.strip()

            if line.startswith("### "):  # H3
                doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith("## "):  # H2
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith("# "):   # H1
                doc.add_heading(line[2:].strip(), level=1)
            elif line.startswith("- "):   # Bullet point
                clean_line = line[2:].strip()

                if "**" in clean_line:
                    parts = clean_line.split("**")
                    paragraph = doc.add_paragraph(style="List Bullet")
                    for i, part in enumerate(parts):
                        if i % 2 == 1:
                            run = paragraph.add_run(part)
                            run.bold = True
                        else:
                            paragraph.add_run(part)
                else:
                    doc.add_paragraph(clean_line, style="List Bullet")

            elif line:
                paragraph = doc.add_paragraph()

                if "**" in line:
                    parts = line.split("**")
                    for i, part in enumerate(parts):
                        if i % 2 == 1:
                            run = paragraph.add_run(part)
                            run.bold = True
                        else:
                            paragraph.add_run(part)
                else:
                    paragraph.add_run(line)
        output_path = f"{userId}_summary.docx"
        doc.save(output_path)
        print("✅ Summary converted to .docx successfully.")
        return output_path

    except Exception as e:
        print("Exception: ",e)


def fetchHistoryService(userId):
    try:
        query="Select filetype,content,llmsummary from summaryhistory where user_id=%s"
        pg=postgresDb()
        result=pg.showData(query,(str(userId),))
        print("history result: ",result)
        return result

    except Exception as e:
        print("Exception",e)
        return {"Exception":e}
    
    finally:
        pg.disconnect()


# pdf, docx -> read, convert bytesIO
# txt, csv -> read, decode, convert StringIO    